"""Генерация демо-данных: справочник + прайсы (XLSX, DOCX) + ZIP-архив.

Запуск:  python scripts/make_samples.py
Создаёт data/samples/ и data/samples/archive.zip для быстрой демонстрации.
"""
from __future__ import annotations

import json
import zipfile
from pathlib import Path

import openpyxl
from docx import Document

OUT = Path(__file__).resolve().parent.parent / "data" / "samples"


def make_catalog():
    services = [
        {"service_name": "Общий анализ крови", "synonyms": ["ОАК", "анализ крови общий"], "category": "лаборатория"},
        {"service_name": "Консультация терапевта", "synonyms": ["приём терапевта"], "category": "консультация"},
        {"service_name": "УЗИ органов брюшной полости", "synonyms": ["УЗИ брюшной полости"], "category": "диагностика"},
        {"service_name": "Биохимический анализ крови", "synonyms": ["биохимия крови"], "category": "лаборатория"},
        {"service_name": "Рентгенография грудной клетки", "synonyms": ["рентген ОГК"], "category": "диагностика"},
    ]
    (OUT / "catalog.json").write_text(json.dumps(services, ensure_ascii=False, indent=2), encoding="utf-8")


def make_xlsx(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Прайс"
    ws.append(["Прайс-лист медицинских услуг", None, None])
    ws.append(["Наименование услуги", "Цена резидент, ₸", "Цена нерезидент, ₸"])
    ws.append(["Общий анализ крови", 2500, 3000])
    ws.append(["Биохимия крови", 7500, 9000])
    ws.append(["УЗИ брюшной полости", 8000, 10000])
    wb.save(path)


def make_docx(path: Path):
    doc = Document()
    doc.add_heading("Прайс клиники Бета", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Услуга", "Цена резидент", "Цена нерезидент"
    for name, r, nr in [
        ("Консультация терапевта", "5000", "6000"),
        ("Рентген ОГК", "4000", "5000"),
        ("ОАК", "2600", "3100"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = name, r, nr
    doc.save(path)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    make_catalog()
    f1 = OUT / "Клиника_Альфа_2025-01-15.xlsx"
    f2 = OUT / "Клиника_Бета_2025-02-01.docx"
    make_xlsx(f1)
    make_docx(f2)
    with zipfile.ZipFile(OUT / "archive.zip", "w") as zf:
        zf.write(f1, f1.name)
        zf.write(f2, f2.name)
    print(f"Готово: {OUT}")
    print("Справочник: catalog.json; Архив: archive.zip")


if __name__ == "__main__":
    main()
