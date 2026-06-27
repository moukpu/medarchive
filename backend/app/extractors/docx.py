"""DOCX extractor: принятие tracked changes + извлечение таблиц.

Принятие отслеживаемых изменений: текст вставок лежит в <w:ins>…<w:t>, текст
удалений — в <w:delText> (другой тег). Поэтому сбор всех <w:t> элементов даёт
ИМЕННО финальную (принятую) версию: вставки попадают, удаления отбрасываются.
"""
from __future__ import annotations

from app.extractors.base import ExtractResult, rows_from_matrix, rows_from_text

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _cell_text(tc) -> str:
    # все <w:t> в ячейке = принятая версия (delText игнорируется автоматически)
    parts = [node.text or "" for node in tc.iter(f"{_W}t")]
    return "".join(parts).strip()


class DocxExtractor:
    def extract(self, path: str) -> ExtractResult:
        from docx import Document

        result = ExtractResult()
        doc = Document(path)
        raw_chunks: list[str] = []

        for table in doc.tables:
            matrix: list[list[str]] = []
            for row in table.rows:
                cells = [_cell_text(cell._tc) for cell in row.cells]
                matrix.append(cells)
                raw_chunks.append("\t".join(cells))
            rows, warnings = rows_from_matrix(matrix)
            result.rows.extend(rows)
            result.warnings.extend(warnings)

        # текст вне таблиц (для raw_content и fallback-парсинга)
        para_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        raw_chunks.append(para_text)
        result.raw_text = "\n".join(raw_chunks)

        if not result.rows and para_text:
            rows, _ = rows_from_text(para_text)
            result.rows.extend(rows)
        if not result.rows:
            result.warnings.append("Не извлечено ни одной позиции из DOCX")
        return result
